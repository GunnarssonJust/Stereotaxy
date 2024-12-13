# -*- coding: utf-8 -*-
"""
Created on Tue Mar  8 15:23:55 2022
Modified for criteria in Helios Clinics Schwerin, Germany on Sat Nov 23-30 2024
.exe-file created with auto-py-to-exe

@authors: Gunnar Just
"""


import numpy as np
import scipy as sp
import tkinter as tk
from tkinter import simpledialog
import os
import docx
from docx2pdf import convert


with open('dvh.txt') as f:
    matrix=[line.split() for line in f]


def search_string_in_file(file_name, string_to_search):
    """Search for the given string in file and return line number"""
    """Count lines numbers from 0"""

    line_number = -1
    # Open the file in read only mode
    with open(file_name, 'r') as read_obj:
        # Read all lines in the file one by one

        for line in read_obj:
            # For each line, check if line contains the string
            line_number += 1
            if string_to_search in line:
                # If yes, then add the line number
                break

    return line_number


# search start line of data

def read_from_line(number_to_start_with, string_to_search):
    with open('dvh.txt') as f:
        for i in np.arange(0, number_to_start_with):
            f, next(f)
            #print(matrix[i])
        for line in f:
            #print(i)
            i = i+1
            if string_to_search in line:
                break
    return i
    



# create numpy-Array from PTV-DVH-Data
def array_from_dvh_data(string_to_search):
    # find length of DVH-Data
    # 1. "Relative dose" and following "Structure:" defines length of data
    start = search_string_in_file('dvh.txt', string_to_search)+26
    #print(start)
    end = read_from_line(start, 'Structure:')
    #print(end)
    len_dvh_data = end-start-2  #-2, for subtract number of empty and start line
    #print(len_dvh_data)
    first_line_dvh = start
    #print(first_line_dvh)
    #print(first_line_dvh)
    dvh_data = np.array(matrix[first_line_dvh:len_dvh_data+first_line_dvh],dtype=float)
    for i in range(0,len_dvh_data):
        dvh_data[i].reshape((1,3))
    
    return(dvh_data)
    

##############################################################################
lastname = matrix[search_string_in_file('dvh.txt', 'Patient Name')][3]
lastname = lastname.replace(",", "")
firstname = matrix[search_string_in_file('dvh.txt', 'Patient Name')][4]
ID = matrix[search_string_in_file('dvh.txt', 'ID')][3]
plan_name_list = matrix[search_string_in_file('dvh.txt', 'Plan')][1:]
separator = ''
plan_name = separator.join(plan_name_list)
coursename = matrix[search_string_in_file('dvh.txt', 'Course')][1]
dosis_prescribed = round(float(matrix[search_string_in_file('dvh.txt','Total dose ')][3]), 2)
dose_prescribed_line = search_string_in_file('dvh.txt','% for dose (%):')

#output of prescribed Dose enclosing PTV(%)
dose_prescribed_percentage = matrix[dose_prescribed_line][4]
#print(dose_prescribed_percentage)
#outputs strived Maximum dose in GTV
dosis_100 = str(float(dosis_verschrieben)/float(dose_prescribed_percentage)/0.01)
#print(dosis_100)

############# Attention, if relative Dose in DVH: data in %, else in Gy 

def find_value(search_word1, search_word2):

    if search_word1 in open('dvh.txt').read():
        if line[0] == 'Relative ':
            D_mean_per = matrix[read_from_line(search_string_in_file('dvh.txt', search_word1)-1, search_word2)][3]
            D_mean_per = round(float(D_mean_per),2)
            D_mean_Gy = str(float(D_mean_per)/100*float(dosis_100))
            D_mean_Gy = round(float(D_mean_Gy), 2)
            
            return(D_mean_Gy, D_mean_per)
        
        else:
            D_mean_Gy = matrix[read_from_line(search_string_in_file('dvh.txt', search_word1)-1, search_word2)][3]
            D_mean_Gy = round(float(D_mean_Gy),2)
            D_mean_per = D_mean_Gy/float(dosis_100)
            D_mean_per = round(D_mean_per*100, 2)  

            return(D_mean_Gy, D_mean_per)

# If there are more than one PTV in DVH, f.e. PTV1, PTV2, etc.:
if 'Structure: PTV ' in open('dvh.txt').read():
    ptv_volume_name = 'Structure: PTV '
else:
    ROOT_2 = tk.Tk()
    ROOT_2.withdraw()
    USER_INPUT = simpledialog.askstring(title='Name of Target volume',prompt = 'The name of target volume is not explicit. \n Please enter the name of the treated volume: ')
    ptv_volume_name = 'Structure: ' +USER_INPUT

if USER_INPUT in open('dvh.txt').read():    
    #type (ptv_volume) = String    
    ptv_volume = matrix[read_from_line(search_string_in_file('dvh.txt', ptv_volume_name)-1, 'Volume')][2]
else: 
    print('Target volume unknown in this plan. Please enter again:')
    ROOT_2 = tk.Tk()
    ROOT_2.withdraw()
    USER_INPUT = simpledialog.askstring(title='Name of Target volume',prompt = 'The name of target volume is not really clear. \n Please enter the name of the treated volume: ')
    ptv_volume_name = 'Structure: ' +USER_INPUT
    ptv_volume = matrix[read_from_line(search_string_in_file('dvh.txt', ptv_volume_name)-1, 'Volume')][2]

def create_DVH_and_abs_rel_doses(structure):
    dvh = array_from_dvh_data(structure)
    
    # If first word in line is "Relative", relative Doses in column 0,
    if line[0] == 'Relative':
        rel_doses = dvh[:,0]
        abs_doses = dvh[:,1]
    # if not, absolute Doses in column 0 and relatives in column 1    
    else:
        rel_doses = dvh[:,1]
        abs_doses = dvh[:,0]
    
    return dvh, abs_doses, rel_doses

#approve, whether relative or absolute dose in column 0:
line_with_relative_dose = read_from_line(search_string_in_file('dvh.txt', 'Relative dose')-1, 'Relative dose')
line = matrix[line_with_relative_dose]




'''
    Create table of PTV-data
'''  
#User input of prescribed percentage of dose
ROOT = tk.Tk()
ROOT.withdraw()
USER_INP = simpledialog.askstring(title='Dose prescription in percent',prompt = 'Please enter prescribed dose in percent(65%, 80% or 100%)')
dose_in_percent = float(USER_INP)

dvh_ptv, abs_doses_ptv, rel_doses_ptv = create_DVH_and_abs_rel_doses(ptv_volume_name)
D_mean_Gy, D_mean_per = find_value(ptv_volume_name, 'Mean')
rel_volumes_ptv = dvh_ptv[:,2]    

def find_D_Vx(x, rel_volumes):
    # calculate shortest difference to x and read dvh_ptv array for absolute und relative Doses at this position:
    D_Vx = round(abs_doses_ptv[np.argmin(abs(rel_volumes-x))], 2) # in Gy
    D_Vx_rel = round(rel_doses_ptv[np.argmin(abs(rel_volumes-x))], 2) # in % der Plandosis
    
    return D_Vx, D_Vx_rel

D_V98, D_V98_rel = find_D_Vx(98, rel_volumes_ptv)
D_V2, D_V2_rel = find_D_Vx(2, rel_volumes_ptv)
D_V50, D_V50_rel = find_D_Vx(50, rel_volumes_ptv)



V_ptv_D100 = dvh_ptv[np.where(rel_doses_ptv==100),2]# result in %
#change type of V_ptv_D100 from array to scalar float
V_ptv_D100 = float(np.take(V_ptv_D100,0))
V_ptv_D100 = float(V_ptv_D100)/100*float(ptv_volume) # result in cm³

# calculate D(V-35mm³) and D(V=35mm³), 35 mm³ according to ICRU91
x = float(ptv_volume) - 0.035 # cm³
y = x/float(ptv_volume) # relative Volume in relation to PTV-Volume

D_Vminus35, D_Vminus35_rel = find_D_Vx(y, rel_volumes_ptv)


z = 0.035/float(ptv_volume) # relative Volume in relation to PTV-Volume
D_Vz, D_Vz_rel = find_D_Vx(z, rel_volumes_ptv)

# create numpy-Array from Corpse/Body-Data
dvh_corpse = array_from_dvh_data('Body')

# prove, where to find relative Doses (row 0 oder 1)
if line[0] == 'Relative dose [%]':
     rel_doses_corpse = dvh_corpse[:,0]
else:
    rel_doses_corpse = dvh_corpse[:,1]

V_iso_D100 = dvh_corpse[np.where(rel_doses_corpse==100), 2]
V_iso_D50 = dvh_corpse[np.where(rel_doses_corpse==50), 2]
V_iso_D100 = float(np.take(V_iso_D100,0))#result in % of Body-Volume
V_iso_D50 = float(np.take(V_iso_D50,0))  #result in % of Body-Volume
print(V_iso_D100)
print(V_iso_D50)
# corpse volume
corpse_volume = matrix[read_from_line(search_string_in_file('dvh.txt', 'Body')-1, 'Volume')][2]

V_iso_D100 = float(V_iso_D100)/100*float(corpse_volume) # result in cm³
V_iso_D50 = float(V_iso_D50)/100*float(corpse_volume) # result in cm³

paddick = float(V_ptv_D100)**2/(float(ptv_volume)*float(V_iso_D100))
#print(corpse_volume)
#print(V_ptv_D100)
#print(ptv_volume)

paddick = round(paddick, 2)

# Gradient-Index = V_iso_D50/V_iso_D100

GI = V_iso_D50/V_iso_D100
GI = round(GI, 2)


###############################################################################
# ORGANS AT RISK   ############################################################
###############################################################################


def findV_Dx(x, dvh, abs_doses, volume, volumePTV):
    V_DxGy = dvh[np.argmin(abs(abs_doses-x)),2] #result in %
    V_DxGy = float(V_DxGy)/100*float(volume) # result in cm³
    if V_DxGy-float(volumePTV) > 0:
        V_DxGy_noPTV = round(V_DxGy-float(volumePTV), 2) # result in cm³ without PTV
    else:
        V_DxGy_noPTV = 0
    
    V_DxGy = round(V_DxGy, 2) # result in cm³ rounded 2 digits right of comma

    return V_DxGy, V_DxGy_noPTV

# brain___________________________________________________________________________

if 'Structure: Brain' in open('dvh.txt').read():

    # # Median - Brain dose
    D_med_Gy_hirn, D_med_per_hirn = find_value('Structure: Brain', 'Median')    
    
    # # Volume of Brain, that receives 10 or 12 Gy
    dvh_Brain = array_from_dvh_data('Structure: Brain')
    
    # # prove, where to find relative Doses (row 0 oder 1)
    if line[0] == 'Relative':
        abs_doses_Brain = dvh_corpse[:,1]
    else:
        abs_doses_Brain = dvh_corpse[:,0]      
        
    # Brain-Volume
    Brain_volume = matrix[read_from_line(search_string_in_file('dvh.txt', 'Structure: Brain')-1, 'Volume')][2]
    
    # #V V10, V12 und V24 inclusive und exclusive PTV Volume    
    V_hirn_D10Gy, V_hirn_D10Gy_noPTV = findV_Dx(10, dvh_Brain, abs_doses_Brain, Brain_volume, ptv_volume)
    V_hirn_D12Gy, V_hirn_D12Gy_noPTV = findV_Dx(12, dvh_Brain, abs_doses_Brain, Brain_volume, ptv_volume)
    V_hirn_D24Gy, V_hirn_D24Gy_noPTV = findV_Dx(24, dvh_Brain, abs_doses_Brain, Brain_volume, ptv_volume)  

# Brain calculations end here________________________________________________________________________________________________________

def output_max(search_structure, structure):
    # structures of risk
    D_max_Gy = find_value(search_structure, 'Max')
    # prove, whether variable has value
    if D_max_Gy:
        D_max_Gy = D_max_Gy[0]
        dvh, abs_doses, rel_doses = create_DVH_and_abs_rel_doses(structure)
        rel_volumes = dvh[:,2]
        D_V2, D_V2_rel = find_D_Vx(2, rel_volumes)
    else:
        D_max_Gy, D_V2, D_V2_rel = None, None, None
    
    return D_max_Gy, D_V2, D_V2_rel

def output_mean(search_structure):
    D_mean = find_value(search_structure, 'Mean')
     # prove, whether variable is assigned
    if D_mean:
        D_mean = D_mean[0]
    else:
        D_mean = None
    
    return D_mean

#TODO: Enter organs at risk of thorax region
##alphabetical order, for loop over array of Structure names, string array by using output_metrics method= output max+output_mean and create output
D_max_Gy_bladder, D_V2_bladder, D_V2_rel_bladder = output_max('Structure: Bladder', 'Bladder')
D_mean_bladder = output_mean('Structure: Bladder')
D_max_Gy_z_bladder, D_V2_z_bladder, D_V2_rel_z_bladder = output_max('Structure: z_Bladder', 'z_Bladder')
D_mean_z_bladder = output_mean('Structure: z_Bladder')
D_mean_bladder = output_mean('Structure: z_Bladder')
D_max_Gy_hs, D_V2_hs, D_V2_rel_hs = output_max('Structure: Brainstem', 'Brainstem')
D_mean_brainstem = output_mean('Structure: Brainstem')
D_max_Gy_Chiasm, D_V2_Chiasm, D_V2_rel_Chiasm = output_max('Structure: Chiasm', 'Chiasm')
D_mean_chiasm = output_mean('Structure: Chiasm')
D_max_Gy_Cochlea_L, D_V2_Cochlea_L, D_V2_rel_Cochlea_L = output_max('Structure: Cochlea_L', 'Cochlea_L')
D_max_Gy_Cochlea_R, D_V2_Cochlea_R, D_V2_rel_Cochlea_R = output_max('Structure: Cochlea_R', 'Cochlea_R')
D_max_Gy_duodenum, D_V2_duodenum, D_V2_rel_duodenum = output_max('Structure: Duodenum', 'Duodenum')
D_mean_duodenum = output_mean('Structure: Duodenum')

D_max_Gy_eye_L, D_V2_eye_L, D_V2_rel_eye_L = output_max('Structure: Eye_L', 'Eye_L')
D_mean_eye_L = output_mean('Structure: Eye_L')
D_max_Gy_eye_R, D_V2_eye_R, D_V2_rel_eye_R = output_max('Structure: Eye_R', 'Eye_R')
D_mean_eye_R = output_mean('Structure: Eye_R')

D_max_Gy_heart, D_V2_heart, D_V2_rel_heart = output_max('Structure: Heart', 'Heart')
D_mean_heart = output_mean('Structure: Heart')

D_max_Gy_intestine, D_V2_intestine, D_V2_rel_intestine = output_max('Structure: Intestine', 'Intestine')
D_mean_intestine = output_mean('Structure: Intestine')
D_max_Gy_kidney_l, D_V2_kidney_l, D_V2_rel_kidney_l = output_max('Structure: Kidney_L', 'Kidney_L')
D_mean_kidney_l = output_mean('Structure: Kidney_L')
D_max_Gy_kidney_r, D_V2_kidney_r, D_V2_rel_kidney_r = output_max('Structure: Kidney_R', 'Kidney_R')
D_mean_kidney_r = output_mean('Structure: Kidney_R')

D_max_Gy_lens_li, D_V2_lens_li, D_V2_rel_lens_li = output_max('Structure: Lens_L', 'Lens_L')
D_mean_lens_li = output_mean('Structure: Lens_L')
D_max_Gy_lens_re, D_V2_lens_re, D_V2_rel_lens_re = output_max('Structure: Lens_R', 'Lens_R')
D_mean_lens_re = output_mean('Structure: Lens_R')
D_max_Gy_liver, D_V2_liver, D_V2_rel_liver = output_max('Structure: Liver', 'Liver')
D_mean_liver = output_mean('Structure: Liver')
D_max_Gy_lung_l, D_V2_lung_l, D_V2_rel_lung_l = output_max('Structure: Lung_L', 'Lung_L')
D_mean_lung_l = output_mean('Structure: Lung_L')
D_max_Gy_lung_r, D_V2_lung_r, D_V2_rel_lung_r = output_max('Structure: Lung_R', 'Lung_R')
D_mean_lung_r = output_mean('Structure: Lung_R')
D_max_Gy_mandible, D_V2_mandible, D_V2_rel_mandible = output_max('Structure: Mandible', 'Mandible')
D_mean_mandible = output_mean('Structure: Mandible')
D_max_Gy_oeso, D_V2_oeso, D_V2_rel_oeso = output_max('Structure: Oesophagus', 'Oesophagus')
D_mean_oeso = output_mean('Structure: Oesophagus')

D_max_Gy_opt_li, D_V2_opt_li, D_V2_rel_opt_li = output_max('Structure: OpticNerve_L', 'OpticNerve_L')
D_mean_opt_li = output_mean('Structure: OpticNerve_L')
D_max_Gy_opt_re, D_V2_opt_re, D_V2_rel_opt_re = output_max('Structure: OpticNerve_R', 'OpticNerve_R')
D_mean_opt_re = output_mean('Structure: OpticNerve_R')
D_max_Gy_oral_cavity, D_V2_oral_cavity, D_V2_rel_oral_cavity = output_max('Structure: Oral_Cavity', 'Oral_Cavity')
D_mean_oral_cavity = output_mean('Structure: Oral_Cavity')
D_max_Gy_pancreas, D_V2_pancreas, D_V2_rel_pancreas = output_max('Structure: Pancreas', 'Pancreas')
D_mean_pancreas = output_mean('Structure: Pancreas')
D_max_Gy_parotid_L, D_V2_parotid_L, D_V2_rel_parotid_L = output_max('Structure: Parotid_L', 'Parotid_L')
D_mean_parotid_L = output_mean('Structure: Parotid_L')
D_max_Gy_parotid_R, D_V2_parotid_R, D_V2_rel_parotid_R = output_max('Structure: Parotid_R', 'Parotid_R')
D_mean_parotid_R = output_mean('Structure: Parotid_R')

D_max_Gy_riva, D_V2_riva, D_V2_rel_riva = output_max('Structure: RIVA', 'RIVA')
D_mean_riva = output_mean('Structure: RIVA')

D_max_Gy_spinalcanal, D_V2_spinalcanal, D_V2_rel_spinalcanal = output_max('Structure: SpinalCanal', 'SpinalCanal')
D_max_Gy_myelon, D_V2_myelon, D_V2_rel_myelon = output_max('Structure: SpinalCord', 'SpinalCord')
D_max_Gy_submandibula_l, D_V2_submandibula_l, D_V2_rel_submandibula_l = output_max('Structure: Submandibula_L', 'Submandibula_L')
D_mean_submandibula_l = output_mean('Structure: Submandibula_L')

D_max_Gy_submandibula_r, D_V2_submandibula_r, D_V2_rel_submandibula_r = output_max('Structure: Submandibula_R', 'Submandibula_R')
D_mean_submandibula_r = output_mean('Structure: Submandibula_R')
#D_max_Gy_ , D_V2_ , D_V2_rel_ = output_max('Structure: ', '')




###############################################################################
# Output generation

from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt

doc = Document()
doc.add_heading('Stereotaxy-Documentation According To ICRU Report 91', 0)

# Patients data

doc.add_heading('Patient data', 3)
para = doc.add_paragraph().add_run()
table = doc.add_table(rows=2, cols=2)

cell = table.cell(0, 0)
cell.text = f"lastname, firstname: "
#cell = table.cell(0,1)
#cell.text = f"{lastname}, {firstname}"

cell = table.cell(0, 1)
cell = cell.paragraphs[0].add_run(f"{lastname}, {firstname}")
cell.bold = True

cell = table.cell(1, 0)
cell.text = f"ID"
cell = table.cell(1,1)
cell.text = f"{ID}"

# Plan data

doc.add_heading('Plan data', 3)
para = doc.add_paragraph().add_run()
table = doc.add_table(rows=4, cols=2)

cell = table.cell(0, 0)
cell.text = f"plan name"
cell = table.cell(0,1)
cell.text = f"{plan_name}"

cell = table.cell(1, 0)
cell.text = f"Course"
cell = table.cell(1,1)
cell.text = f"{coursename}"

cell = table.cell(2, 0)
cell.text = f"Planned-Dose ("+ str(dose_in_percent) +" %)"
cell = table.cell(2,1)
cell.text = f"{dosis_verschrieben} Gy"

cell = table.cell(3, 0)
cell.text = f"Maximum dose (100 %)"
cell = table.cell(3,1)
#print(type(dosis_100))
#print(type(dose_in_percent))
cell.text = f"{round(float(dosis_100)/dose_in_percent*100,2)} Gy"

# PTV- Evaluation

doc.add_heading('PTV', 3)
# Volumen, Paddick, Gradient, Dnearminx2, Dnearmaxx2, D50, D mean

para = doc.add_paragraph().add_run()
table = doc.add_table(rows=9, cols=2)

cell = table.cell(0, 0)
cell.text = f"Volume"
cell = table.cell(0,1)
cell.text = f"{ptv_volume} cm³"

cell = table.cell(1, 0)
cell.text = f"Paddick Conformity Index *"


cell = table.cell(1,1)
cell = cell.paragraphs[0].add_run(f"{paddick}")
cell.bold = True
font = cell.font
if paddick <= 0.5:    
    font.color.rgb = RGBColor(139, 0, 0)
elif 0.5 < paddick < 0.7:
    font.color.rgb = RGBColor(255, 165, 0)
elif paddick >= 0.7:
    font.color.rgb = RGBColor(0, 128, 0)



cell = table.cell(2, 0)
cell.text = f"Gradient Index *"

cell = table.cell(2,1)
#cell.text = f"{GI}"
cell = cell.paragraphs[0].add_run(f"{GI}")
cell.bold = True
font = cell.font
if GI <= 4:    
    font.color.rgb = RGBColor(0, 128, 0)
elif 4 < GI < 6:
    font.color.rgb = RGBColor(255, 165, 0)
elif GI >= 6:
    font.color.rgb = RGBColor(139, 0, 0)


cell = table.cell(3, 0)
cell = cell.paragraphs[0].add_run('D')
cell = table.cell(3, 0)
cell = cell.paragraphs[0].add_run('near-min')
cell.font.subscript = True
cell = table.cell(3, 0)
cell = cell.paragraphs[0].add_run('(V=98 %)')

cell = table.cell(3, 1)
cell.text = f"{D_V98} Gy bzw. {D_V98_rel} % of Planned-Dose"

cell = table.cell(4, 0)
cell = cell.paragraphs[0].add_run('D')
cell = table.cell(4, 0)
cell = cell.paragraphs[0].add_run('near-max')
cell.font.subscript = True
cell = table.cell(4, 0)
cell = cell.paragraphs[0].add_run('(V=2 %)')

cell = table.cell(4, 1)
cell.text = f"{D_V2} Gy bzw. {D_V2_rel} % of Planned-Dose"

cell = table.cell(5, 0)
cell = cell.paragraphs[0].add_run('D')
cell = table.cell(5, 0)
cell = cell.paragraphs[0].add_run('near-min')
cell.font.subscript = True
cell = table.cell(5, 0)
cell = cell.paragraphs[0].add_run('(V-35 mm³)')

cell = table.cell(5, 1)
cell.text = f"{D_Vminus35} Gy bzw. {D_Vminus35_rel} % of Planned-Dose"

cell = table.cell(6, 0)
cell = cell.paragraphs[0].add_run('D')
cell = table.cell(6, 0)
cell = cell.paragraphs[0].add_run('near-max')
cell.font.subscript = True
cell = table.cell(6, 0)
cell = cell.paragraphs[0].add_run('(V=35 mm³)')

cell = table.cell(6, 1)
cell.text = f"{D_Vz} Gy bzw. {D_Vz_rel} % of Planned-Dose"

cell = table.cell(7, 0)
cell = cell.paragraphs[0].add_run('D(V=50 %)')

cell = table.cell(7, 1)
cell.text = f"{D_V50} Gy bzw. {D_V50_rel} % of Planned-Dose"

cell = table.cell(8, 0)
cell = cell.paragraphs[0].add_run('D')
cell = table.cell(8, 0)
cell = cell.paragraphs[0].add_run('mean')
cell.font.subscript = True

cell = table.cell(8, 1)
cell.text = f"{D_mean_Gy} Gy bzw. {D_mean_per} % of Planned-Dose"


para = doc.add_paragraph().add_run()
para = doc.add_paragraph().add_run()

##############################################################################
# # # # # # # # # If Organs at Risk should be evaluated: delete comment# 
#Organs at risk at next page, Open Office gives a changed Layout, Format than Word!!!
# doc.add_page_break()
doc.add_heading('Risikoorgane', 2)

# Brain-Evaluation

if 'Structure: Brain' in open('dvh.txt').read():

    doc.add_heading('Brain', 3)
    para = doc.add_paragraph().add_run()
    table = doc.add_table(rows=7, cols=2)
    
    cell = table.cell(0, 0)
    cell = cell.paragraphs[0].add_run('V')
    cell = table.cell(0, 0)
    cell = cell.paragraphs[0].add_run('Brain-PTV')
    cell.font.subscript = True
    cell = table.cell(0, 0)
    cell = cell.paragraphs[0].add_run('(10 Gy)')
    
    cell = table.cell(0, 1)
    cell.text = f"{V_hirn_D10Gy_noPTV} cm³"
    
    
    cell = table.cell(1, 0)
    cell = cell.paragraphs[0].add_run('V')
    cell = table.cell(1, 0)
    cell = cell.paragraphs[0].add_run('Brain-PTV')
    cell.font.subscript = True
    cell = table.cell(1, 0)
    cell = cell.paragraphs[0].add_run('(12 Gy)')
    
    cell = table.cell(1, 1)
    cell.text = f"{V_hirn_D12Gy_noPTV} cm³"
    
    
    cell = table.cell(2, 0)
    cell = cell.paragraphs[0].add_run('V')
    cell = table.cell(2, 0)
    cell = cell.paragraphs[0].add_run('Brain-PTV')
    cell.font.subscript = True
    cell = table.cell(2, 0)
    cell = cell.paragraphs[0].add_run('(24 Gy)')
    
    cell = table.cell(2, 1)
    cell.text = f"{V_hirn_D24Gy_noPTV} cm³"
    
    
    cell = table.cell(3, 0)
    cell = cell.paragraphs[0].add_run('V')
    cell = table.cell(3, 0)
    cell = cell.paragraphs[0].add_run('Brain inkl. PTV')
    cell.font.subscript = True
    cell = table.cell(3, 0)
    cell = cell.paragraphs[0].add_run('(10 Gy)')
    
    cell = table.cell(3, 1)
    cell.text = f"{V_hirn_D10Gy} cm³"
    
    
    cell = table.cell(4, 0)
    cell = cell.paragraphs[0].add_run('V')
    cell = table.cell(4, 0)
    cell = cell.paragraphs[0].add_run('Brain inkl. PTV')
    cell.font.subscript = True
    cell = table.cell(4, 0)
    cell = cell.paragraphs[0].add_run('(12 Gy)')
    
    cell = table.cell(4, 1)
    cell.text = f"{V_hirn_D12Gy} cm³"
    
    cell = table.cell(5, 0)
    cell = cell.paragraphs[0].add_run('V')
    cell = table.cell(5, 0)
    cell = cell.paragraphs[0].add_run('Brain inkl. PTV')
    cell.font.subscript = True
    cell = table.cell(5, 0)
    cell = cell.paragraphs[0].add_run('(24 Gy)')
    
    cell = table.cell(5, 1)
    cell.text = f"{V_hirn_D24Gy} cm³"
    
    
    cell = table.cell(6, 0)
    cell = cell.paragraphs[0].add_run('D')
    cell = table.cell(6, 0)
    cell = cell.paragraphs[0].add_run('median')
    cell.font.subscript = True
    
    cell = table.cell(6, 1)
    cell.text = f"{D_med_Gy_hirn} Gy"


        
def create_output(search_word, D_max=False, D_V2=False, D_mean=False):
    
    
    if search_word in open('dvh.txt').read():

        doc.add_heading(search_word, 3)
        para = doc.add_paragraph().add_run()

        if (D_max and D_V2 and D_mean):
            table = doc.add_table(rows=3, cols=2)
            
            cell = table.cell(0, 0)
            cell = cell.paragraphs[0].add_run('D')
            cell = table.cell(0, 0)
            cell = cell.paragraphs[0].add_run('max')
            cell.font.subscript = True
        
            cell = table.cell(0, 1)
            cell.text = f"{D_max} Gy"
        
            cell = table.cell(1, 0)
            cell = cell.paragraphs[0].add_run('D')
            cell = table.cell(1, 0)
            cell = cell.paragraphs[0].add_run('near-max')
            cell.font.subscript = True
            cell = table.cell(1, 0)
            cell = cell.paragraphs[0].add_run('(V=2 %)')
        
            cell = table.cell(1, 1)
            cell.text = f"{D_V2} Gy"
            
            cell = table.cell(2, 0)
            cell = cell.paragraphs[0].add_run('D')
            cell = table.cell(2, 0)
            cell = cell.paragraphs[0].add_run('mean')
            cell.font.subscript = True
        
            cell = table.cell(2, 1)
            cell.text = f"{D_mean} Gy"
            
        elif (D_max and D_V2):
            table = doc.add_table(rows=2, cols=2)
            
            cell = table.cell(0, 0)
            cell = cell.paragraphs[0].add_run('D')
            cell = table.cell(0, 0)
            cell = cell.paragraphs[0].add_run('max')
            cell.font.subscript = True
        
            cell = table.cell(0, 1)
            cell.text = f"{D_max} Gy"
        
        
            cell = table.cell(1, 0)
            cell = cell.paragraphs[0].add_run('D')
            cell = table.cell(1, 0)
            cell = cell.paragraphs[0].add_run('near-max')
            cell.font.subscript = True
            cell = table.cell(1, 0)
            cell = cell.paragraphs[0].add_run('(V=2 %)')
        
            cell = table.cell(1, 1)
            cell.text = f"{D_V2} Gy"
            
        elif (D_mean):
            table = doc.add_table(rows=1, cols=2)
            
            cell = table.cell(0, 0)
            cell = cell.paragraphs[0].add_run('D')
            cell = table.cell(0, 0)
            cell = cell.paragraphs[0].add_run('mean')
            cell.font.subscript = True
        
            cell = table.cell(0, 1)
            cell.text = f"{D_mean} Gy"
    else: print('Structure ' +search_word+' was not chosen for evaluation.')  

create_output(' Bladder', D_max_Gy_bladder, D_V2_bladder, D_mean_bladder)
create_output(' z_Bladder', D_max_Gy_z_bladder, D_V2_z_bladder, D_mean_z_bladder)
create_output(' Brainstem', D_max_Gy_hs, D_V2_hs, D_mean_brainstem)
create_output(' Chiasm', D_max_Gy_Chiasm, D_V2_Chiasm, D_mean_chiasm)
create_output(' Cochlea_L', D_max_Gy_Cochlea_L, D_V2_Cochlea_L)
create_output(' Cochlea_R', D_max_Gy_Cochlea_R, D_V2_Cochlea_R)
create_output(' Duodenum', D_max_Gy_duodenum, D_V2_duodenum,D_mean_duodenum)
create_output(' Heart', D_max_Gy_heart, D_V2_heart, D_mean_heart)
create_output(' Intestine', D_max_Gy_intestine, D_V2_intestine, D_mean_intestine)
create_output(' Kidney_L', D_max_Gy_kidney_l, D_V2_kidney_l, D_mean_kidney_l)
create_output(' Kidney_R', D_max_Gy_kidney_r, D_V2_kidney_r, D_mean_kidney_r)
create_output(' Lens_L', D_max_Gy_lens_li, D_V2_lens_li, D_mean_lens_li)
create_output(' Lens_R', D_max_Gy_lens_re, D_V2_lens_re, D_mean_lens_re)
create_output(' Liver', D_max_Gy_liver, D_V2_liver, D_mean_liver)
create_output(' Lung_L', D_max_Gy_lung_l, D_V2_lung_l, D_mean_lung_l)
create_output(' Lung_R', D_max_Gy_lung_r, D_V2_lung_r, D_mean_lung_r)
create_output(' Mandible', D_max_Gy_mandible, D_V2_mandible, D_mean_mandible)

create_output(' Oesophagus', D_max_Gy_oeso, D_V2_oeso, D_mean_oeso)
create_output(' OpticNerve_L', D_max_Gy_opt_li, D_V2_opt_li, D_mean_opt_li)
create_output(' OpticNerve_R', D_max_Gy_opt_re, D_V2_opt_re, D_mean_opt_re)
create_output(' Oral_Cavity', D_max_Gy_oral_cavity, D_V2_oral_cavity, D_mean_oral_cavity)
create_output(' Pancreas', D_max_Gy_pancreas, D_V2_pancreas, D_mean_pancreas)
create_output(' Parotid_L', D_max_Gy_parotid_L, D_V2_parotid_L, D_mean_parotid_L)
create_output(' Parotid_R', D_max_Gy_parotid_R, D_V2_parotid_R, D_mean_parotid_R)
create_output(' RIVA', D_max_Gy_riva, D_V2_riva, D_mean_riva)
create_output(' SpinalCanal', D_max_Gy_spinalcanal, D_V2_spinalcanal)
create_output(' SpinalCord', D_max_Gy_myelon, D_V2_myelon)
create_output(' Submandibula_L', D_max_Gy_submandibula_l, D_V2_submandibula_l, D_mean_submandibula_l)
create_output(' Submandibula_R', D_max_Gy_submandibula_r, D_V2_submandibula_r, D_mean_submandibula_r)

# Add footnote
# >= in Unicode
s = u'\u2265'
# <= in Unicode
t = u'\u2264'
section = doc.sections[0] 
footer = section.footer

#footer.is_linked_to_previous = True
#first_page_footer 
footer_para = footer.paragraphs[0].add_run('* Calculation after ICRU report 91 on prescribing, recording, and reporting of'+
' stereotactic treatments with small photon beams - Statement from the DEGRO/DGMP working group stereotactic radiotherapy and surgery (2019), with Paddick = 1/CI')
font = footer_para.font
font.size = Pt(7)
footer_para.add_break()
footer_para = footer.paragraphs[0].add_run(f'Paddick Conformity Index {t} 0.5 rot, 0.5 < PCI < 0.7 gelb, {s} 0.7 grün')
font = footer_para.font
font.size = Pt(7)
footer_para.add_break()
footer_para = footer.paragraphs[0].add_run(f'Gradient Index {t} 4 grün, 4 < GI < 6 gelb, {s} 6 rot')
font = footer_para.font
font.size = Pt(7)



# delete dvh.txt
#os.remove("dvh.txt")
#print('file dvh.txt deleted successfully')


import time
today = time.strftime("%Y%m%d-%H%M%S")

# Now save the document to a location 
# doc.save(f'{lastname}, {firstname} ({ID})_{today}.docx')


# Now save the document to a location 
path = f'C:/Users/Your_username/Desktop//{lastname}, {firstname} ({ID})'
#or save wherever You want

if os.path.isdir(path):
    print('Directory still exists.')
    doc.save(f'{path}/{lastname}, {firstname} ({ID})_{today}.docx')
#creates a .pdf-document from .docx-Protocol
    convert(f'{path}/{lastname}, {firstname} ({ID})_{today}.docx')#,OtherFolder(AC_Export?\{lastname}, {firstname} ({ID})_{today}.pdf')
else:
    print('Directory created honestly.')
    os.makedirs(path)
    doc.save(f'{path}/{lastname}, {firstname} ({ID})_{today}.docx')
#erstellt ein .pdf-Dokument aus dem .docx-Protokoll
    convert(f'{path}/{lastname}, {firstname} ({ID})_{today}.docx')#,OtherFolder\{lastname}, {firstname} ({ID})_{today}.pdf')

print(f'Creating report was successful!\nSaved to C:/Users/username/Desktop/{lastname}, {firstname} ({ID})')
#or path of Your choice
input("Press Enter key to finish .py program...")





